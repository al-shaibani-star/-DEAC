# -*- coding: utf-8 -*-
"""Build 09_algorithms.docx"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BLUE = RGBColor(0,101,189)
GREEN = RGBColor(34,139,34)
GRAY = RGBColor(136,136,136)
RED = RGBColor(204,51,51)
doc = Document()
for s in doc.sections:
    s.top_margin=Cm(2); s.bottom_margin=Cm(2); s.left_margin=Cm(2); s.right_margin=Cm(2)

def hdr(n,t):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(14)
    p.paragraph_format.space_after=Pt(6)
    sh=p._element.get_or_add_pPr()
    sd=sh.makeelement(qn('w:shd'),{qn('w:fill'):'D5E8F0',qn('w:val'):'clear'})
    sh.append(sd)
    r=p.add_run(f'Algorithm {n}: '); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE
    r=p.add_run(t); r.bold=True; r.font.size=Pt(12)

def io(l,t):
    p=doc.add_paragraph()
    p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run(f'{l}: '); r.bold=True; r.font.size=Pt(9)
    r=p.add_run(t); r.font.name='Cambria Math'; r.font.size=Pt(9); r.italic=True

def ln(n,ind,parts):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(1)
    p.paragraph_format.space_after=Pt(1)
    r=p.add_run(f'{n:>2}  '); r.font.name='Consolas'; r.font.size=Pt(8.5); r.font.color.rgb=GRAY
    if ind>0:
        r=p.add_run('    '*ind); r.font.name='Consolas'; r.font.size=Pt(8.5)
    for x in parts:
        if isinstance(x,str):
            r=p.add_run(x); r.font.name='Consolas'; r.font.size=Pt(8.5)
        elif 'kw' in x:
            r=p.add_run(x['kw']); r.font.name='Consolas'; r.font.size=Pt(8.5); r.bold=True; r.font.color.rgb=BLUE
        elif 'cm' in x:
            r=p.add_run('  // '+x['cm']); r.font.name='Consolas'; r.font.size=Pt(7.5); r.italic=True; r.font.color.rgb=GREEN
        elif 'ref' in x:
            r=p.add_run(x['ref']); r.font.name='Consolas'; r.font.size=Pt(7.5); r.font.color.rgb=RED
        elif 'eq' in x:
            r=p.add_run(x['eq']); r.font.name='Cambria Math'; r.font.size=Pt(9); r.italic=True
        elif 'fn' in x:
            r=p.add_run(x['fn']); r.font.name='Consolas'; r.font.size=Pt(8.5); r.bold=True

# Title
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Pseudocode: GWO-DE-NM Hybrid Graph Clustering')
r.bold=True; r.font.size=Pt(16); r.font.color.rgb=BLUE
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(16)
r=p.add_run('A Hybrid Clustering Algorithm Based on Differential Evolution\nfor High-Dimensional Data Processing')
r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY

# ALG 1: Main
hdr(1,'GWO\u2013DE\u2013NM Hybrid Graph Clustering (Main)')
io('Input','Data matrix X \u2208 \u211D^(n\u00D7d), seeds S')
io('Output','Cluster labels \u0177 \u2208 {1,...,C}')
ln(1,0,[{'kw':'Define '},'search bounds: ',{'eq':'\u03B8 = (k, d\u2032, q, r)'}])
ln(2,0,[{'eq':'lb=(5,5,0,0.01)'},', ',{'eq':'ub=(80,100,0.40,0.90)'}])
ln(3,0,[{'kw':'FOR '},'each seed ',{'eq':'s\u2208S'},{'kw':' DO'}])
ln(4,1,[{'eq':'\u0177_B'},' = ',{'fn':'GraphCluster'},'(X, k=35, d\u2032=40, q=0.10, r=0.25)',{'ref':'  [Alg.6]'}])
ln(5,1,[{'eq':'\u03B8*_gwo'},', ',{'eq':'\u210B'},' = ',{'fn':'LevGWO'},'(X, lb, ub, s)',{'ref':'  [Alg.2]'}])
ln(6,1,[{'eq':'\u03B8*_de'},' = ',{'fn':'WarmDE'},'(X, ',{'eq':'\u210B'},', lb, ub, s)',{'ref':'  [Alg.4]'}])
ln(7,1,[{'eq':'\u03B8*'},' = ',{'fn':'NelderMead'},'(X, ',{'eq':'\u03B8*_de'},')',{'ref':'  [Alg.5]'}])
ln(8,1,[{'eq':'\u0177_H'},' = ',{'fn':'GraphCluster'},'(X, ',{'eq':'\u03B8*'},')',{'ref':'  [Alg.6]'}])
ln(9,1,[{'kw':'IF '},{'fn':'QualityGate'},'(',{'eq':'\u0177_B'},', ',{'eq':'\u0177_H'},', n)',{'kw':' THEN'},{'ref':'  [Alg.8]'}])
ln(10,2,[{'eq':'\u0177_s'},' = ',{'eq':'\u0177_B'},{'cm':'revert to baseline'}])
ln(11,1,[{'kw':'ELSE'}])
ln(12,2,[{'eq':'\u0177_s'},' = ',{'eq':'\u0177_H'},{'cm':'accept hybrid'}])
ln(13,1,[{'kw':'END IF'}])
ln(14,0,[{'kw':'END FOR'}])
ln(15,0,[{'kw':'RETURN '},'best ',{'eq':'\u0177_s'},' across seeds'])

# ALG 2: LevGWO
doc.add_page_break()
hdr(2,'LevGWO \u2014 L\u00E9vy-Flight Grey Wolf Optimizer')
io('Input','Data X, bounds [lb,ub], seed s')
io('Output','Best \u03B8*_gwo, elite set \u210B')
ln(1,0,['W=12, T=12, ',{'eq':'\u03B2=1.5'}])
ln(2,0,[{'kw':'Initialize '},{'eq':'{x_i}'},' uniformly in ',{'eq':'[lb,ub]\u2074'}])
ln(3,0,['Evaluate ',{'eq':'f(x_i)'},' \u2200i',{'ref':'  [Alg.7]'}])
ln(4,0,[{'kw':'FOR '},'t=1 ',{'kw':'TO '},'T ',{'kw':'DO'}])
ln(5,1,['Rank: ',{'eq':'\u03B1,\u03B2,\u03B4'},' = top-3 wolves'])
ln(6,1,[{'eq':'a=2(1\u2212t/T)'},{'cm':'linearly decreasing'}])
ln(7,1,[{'eq':'w_L=0.5(1\u2212t/T)+0.01'},{'cm':'L\u00E9vy weight decay'}])
ln(8,1,[{'kw':'FOR '},'i=1 ',{'kw':'TO '},'W ',{'kw':'DO'}])
ln(9,2,[{'kw':'FOR '},'j=1 ',{'kw':'TO '},'4 ',{'kw':'DO'},{'cm':'GWO position update'}])
ln(10,3,[{'eq':'X1=\u03B1_j\u2212A\u2081|2r\u2082\u03B1_j\u2212x_{i,j}|'}])
ln(11,3,[{'eq':'X2=\u03B2_j\u2212A\u2082|2r\u2082\u03B2_j\u2212x_{i,j}|'}])
ln(12,3,[{'eq':'X3=\u03B4_j\u2212A\u2083|2r\u2082\u03B4_j\u2212x_{i,j}|'}])
ln(13,3,[{'eq':'x_{i,j}=(X1+X2+X3)/3'}])
ln(14,2,[{'kw':'END FOR'}])
ln(15,2,[{'eq':'L'},' = ',{'fn':'MantegnaLevy'},'(4,',{'eq':'\u03B2'},')',{'ref':'  [Alg.3]'}])
ln(16,2,[{'eq':'x_i = x_i + w_L\u00B7L\u00B7(x_i\u2212\u03B1)'},{'cm':'L\u00E9vy perturbation'}])
ln(17,2,['clamp(',{'eq':'x_i'},', lb, ub)'])
ln(18,1,[{'kw':'END FOR'}])
ln(19,1,['Re-evaluate ',{'eq':'f(x_i)'},' \u2200i'])
ln(20,0,[{'kw':'END FOR'}])
ln(21,0,[{'eq':'\u210B'},' = top-5 solutions',{'cm':'elite set for DE'}])
ln(22,0,[{'kw':'RETURN '},{'eq':'\u03B8*_gwo=\u03B1'},', ',{'eq':'\u210B'}])

# ALG 3: Mantegna
hdr(3,"Mantegna\u2019s L\u00E9vy Flight Step")
io('Input','Dimension D, \u03B2=1.5')
io('Output','L\u00E9vy step L \u2208 \u211D\u1D30')
ln(1,0,[{'eq':'\u03C3_u = (\u0393(1+\u03B2)\u00B7sin(\u03C0\u03B2/2) / (\u0393((1+\u03B2)/2)\u00B7\u03B2\u00B72^((\u03B2-1)/2)))^(1/\u03B2)'}])
ln(2,0,[{'eq':'u ~ N(0, \u03C3_u\u00B2\u00B7I_D)'},',  ',{'eq':'v ~ N(0, I_D)'}])
ln(3,0,[{'eq':'L = u / |v|^(1/\u03B2)'},{'cm':'heavy-tailed jumps'}])
ln(4,0,[{'kw':'RETURN '},{'eq':'L'}])

# ALG 4: WarmDE
doc.add_page_break()
hdr(4,'WarmDE \u2014 Warm-Started Differential Evolution')
io('Input','Data X, GWO elite \u210B, bounds [lb,ub], seed s')
io('Output','Optimized \u03B8*_de')
ln(1,0,['G=30, P=12\u00D74=48, ',{'eq':'|\u210B|=5'},{'cm':'elite from GWO'}])
ln(2,0,['pop = [',{'eq':'\u210B\u2081,...,\u210B\u2085'},', random\u2081,...,random\u2084\u2083]',{'cm':'warm-start'}])
ln(3,0,[{'eq':'\u03B1_gwo=\u210B\u2081'},', ',{'eq':'f_\u03B1=f(\u03B1_gwo)'},{'cm':'elitist ref'}])
ln(4,0,[{'kw':'FOR '},'g=1 ',{'kw':'TO '},'G ',{'kw':'DO'}])
ln(5,1,[{'kw':'FOR '},'i=1 ',{'kw':'TO '},'P ',{'kw':'DO'}])
ln(6,2,[{'eq':'v_i=x_{r1}+F\u00B7(x_{r2}\u2212x_{r3})'},{'cm':'mutation'}])
ln(7,2,[{'eq':'u_i'},'=crossover(',{'eq':'x_i,v_i'},',CR)'])
ln(8,2,[{'kw':'IF '},{'eq':'f(u_i)<f(x_i)'},{'kw':' THEN '},{'eq':'x_i=u_i'},{'kw':' END IF'}])
ln(9,1,[{'kw':'END FOR'}])
ln(10,0,[{'kw':'END FOR'}])
ln(11,0,[{'eq':'\u03B8*_de=argmin_i f(x_i)'}])
ln(12,0,[{'kw':'IF '},{'eq':'f_\u03B1<f(\u03B8*_de)'},{'kw':' THEN '},{'eq':'\u03B8*_de=\u03B1_gwo'},{'kw':' END IF'},{'cm':'elitist guarantee'}])
ln(13,0,[{'kw':'RETURN '},{'eq':'\u03B8*_de'}])

# ALG 5: NM
hdr(5,'Nelder\u2013Mead Local Refinement')
io('Input','Initial \u03B8*_de')
io('Output','Refined \u03B8*')
ln(1,0,['M=50, simplex ',{'eq':'S={\u03B8*_de + perturbations}'}])
ln(2,0,[{'kw':'FOR '},'m=1 ',{'kw':'TO '},'M ',{'kw':'DO'}])
ln(3,1,['Sort: ',{'eq':'f(x\u2081)\u2264...\u2264f(x\u2085)'}])
ln(4,1,['Centroid: ',{'eq':'x\u0305=(1/4)\u2211x_i (excl worst)'}])
ln(5,1,['Reflect \u2192 Expand \u2192 Contract \u2192 Shrink'])
ln(6,0,[{'kw':'END FOR'}])
ln(7,0,[{'eq':'\u03B8*=x\u2081'},{'cm':'best vertex'}])
ln(8,0,[{'kw':'IF '},{'eq':'f(\u03B8*)>f(\u03B8*_de)'},{'kw':' THEN '},{'eq':'\u03B8*=\u03B8*_de'},{'kw':' END IF'}])
ln(9,0,[{'kw':'RETURN '},{'eq':'\u03B8*'}])

# ALG 6: Pipeline
doc.add_page_break()
hdr(6,'GraphCluster \u2014 Graph-Based Clustering Pipeline')
io('Input','X \u2208 \u211D^(n\u00D7d), \u03B8=(k,d\u2032,q,r)')
io('Output','Labels \u0177, metrics (C,Q,S)')
ln(1,0,[{'eq':'Z=PCA(Scale(X),d\u2032)'},{'cm':'Step 1: dim reduction'}])
ln(2,0,[{'eq':'N_i'},'=k-NN(',{'eq':'z_i'},', cosine)',{'cm':'Step 2: KNN graph'}])
ln(3,0,[{'eq':'w_{ij}=(1+cos(z_i,z_j))/2'},{'cm':'similarity to weight'}])
ln(4,0,[{'eq':'w_{ij}=(w_{ij}+w_{ji})/2'},{'cm':'Step 3: symmetrize'}])
ln(5,0,['Remove weak edges ',{'eq':'w<Q_q'},{'cm':'Step 4: prune'}])
ln(6,0,['Bridge disconnected components',{'cm':'Step 5: connect'}])
ln(7,0,[{'eq':'\u0177'},'=Louvain(G, resolution=r)',{'cm':'Step 6: cluster'}])
ln(8,0,['Merge tiny clusters (<0.005n)',{'cm':'Step 7: post-process'}])
ln(9,0,[{'eq':'Q'},'=modularity, ',{'eq':'S'},'=silhouette',{'cm':'Step 8: metrics'}])
ln(10,0,[{'kw':'RETURN '},{'eq':'\u0177, C, Q, S'}])

# ALG 7: Fitness
hdr(7,'Composite Fitness Function')
io('Input','\u03B8=(k,d\u2032,q,r), subsample X\u2032\u2286X')
io('Output','f(\u03B8) \u2014 lower is better')
ln(1,0,[{'eq':'\u0177,C,Q,S'},' = ',{'fn':'GraphCluster'},'(X\u2032,\u03B8)'])
ln(2,0,[{'kw':'IF '},{'eq':'C<2'},{'kw':' RETURN '},{'eq':'+\u221E'},{'kw':' END IF'}])
ln(3,0,[{'eq':'MOD\u2080\u2081=(Q+1)/2'},', ',{'eq':'SIL\u2080\u2081=(S+1)/2'},', ',{'eq':'BAL\u2080\u2081=H(\u0177)/log(C)'}])
ln(4,0,[{'eq':'score = 0.50\u00B7MOD + 0.40\u00B7SIL + 0.10\u00B7BAL'}])
ln(5,0,['Penalties: fragmentation + singleton + tiny + disconnect'])
ln(6,0,[{'kw':'RETURN '},{'eq':'\u2212(score\u2212p)'}])

# ALG 8: QualityGate
hdr(8,'QualityGate \u2014 3-Layer Adaptive Protection')
io('Input','Baseline B, Hybrid H, size n')
io('Output','TRUE=baseline, FALSE=hybrid')
ln(1,0,[{'kw':'IF '},{'eq':'C_B<2'},{'kw':' RETURN FALSE'},{'cm':'degenerate \u2192 accept'}])
ln(2,0,[{'eq':'s_B=0.40\u00B7MOD+0.35\u00B7SIL+0.25\u00B7BAL'},{'cm':'quality scores'}])
ln(3,0,[{'eq':'\u03C1=max(4, \u221A(n/50))'},{'cm':'adaptive threshold'}])
ln(4,0,[{'kw':'IF '},{'eq':'C_H\u2265\u03C1\u00B7C_B'},{'kw':' RETURN TRUE'},{'cm':'L1: extreme'}])
ln(5,0,[{'kw':'IF '},{'eq':'C_H>(\u03C1/2)\u00B7C_B AND \u0394s<0.06'},{'kw':' RETURN TRUE'},{'cm':'L2: moderate'}])
ln(6,0,[{'kw':'IF '},{'eq':'s_B>s_H'},{'kw':' RETURN TRUE'},{'cm':'L3: quality'}])
ln(7,0,[{'kw':'RETURN FALSE'},{'cm':'all passed'}])

# Call Hierarchy
p=doc.add_paragraph()
p.paragraph_format.space_before=Pt(16)
r=p.add_run('Call Hierarchy'); r.bold=True; r.font.size=Pt(14); r.font.color.rgb=BLUE
for h in [
    'Alg.1 (Main)',
    '  \u251C Alg.6 (GraphCluster) \u2190 Baseline',
    '  \u251C Alg.2 (LevGWO) \u2190 Phase 1',
    '  \u2502   \u251C Alg.3 (Mantegna)',
    '  \u2502   \u2514 Alg.7 (Fitness) \u2192 Alg.6',
    '  \u251C Alg.4 (WarmDE) \u2190 Phase 2',
    '  \u251C Alg.5 (Nelder-Mead) \u2190 Phase 3',
    '  \u251C Alg.6 (GraphCluster) \u2190 Hybrid eval',
    '  \u2514 Alg.8 (QualityGate) \u2190 Protection',
]:
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(1)
    p.paragraph_format.space_after=Pt(1)
    r=p.add_run(h); r.font.name='Consolas'; r.font.size=Pt(9)

out = r'E:\Strudin\Mo01\GWO_DE_NM_Hybrid\paper\algorithms\09_algorithms.docx'
doc.save(out)
print(f'SUCCESS: {out}')
